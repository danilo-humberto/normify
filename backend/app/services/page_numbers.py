from __future__ import annotations

"""
page_numbers.py — Inserts ABNT-compliant page numbering.

NBR 14724 rules:
  - Arabic numerals
  - Top-right corner of the page
  - Counted from the first page of the document (cover = page 1)
  - Numbers only appear from the first textual element (Introduction)
  - Cover and pre-textual pages are counted but not numbered
"""

from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def add_page_numbers(doc: DocumentObject, font_name: str, start_numbering_from: int = 2) -> None:
    """
    Add page numbers to the header of all sections.

    Args:
        doc: The document to modify.
        font_name: Font to use for page numbers.
        start_numbering_from: Physical page number where the number first appears.
                              Default 2 means the cover (page 1) has no number shown.
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True
        _clear_header(section.header)
        _clear_header(section.first_page_header)
        _add_number_to_header(section.header, font_name)
        # first_page_header stays empty → cover has no number


def _clear_header(header: object) -> None:
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            run.text = ""


def _add_number_to_header(header: object, font_name: str) -> None:
    """Insert a right-aligned field code {PAGE} into the header."""
    if not header.paragraphs:
        paragraph = header.add_paragraph()
    else:
        paragraph = header.paragraphs[0]

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(10)

    # Insert Word field: { PAGE }
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
