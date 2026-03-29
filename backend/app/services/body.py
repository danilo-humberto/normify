from __future__ import annotations

"""
body.py — Applies ABNT NBR 14724 formatting to the document body.

Responsibilities:
  - Detect and remove an existing cover (when has_existing_cover=True)
  - Apply body paragraph styles (font, size, spacing, indent)
  - Apply heading styles to section titles
  - Apply references styles to reference entries
  - Insert page breaks before each major section heading
"""

import logging

from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from app.schemas.document import CoverData
from app.utils.docx_helpers import (
    apply_run_style,
    is_references_heading,
    is_section_heading,
    paragraph_has_page_break,
)

logger = logging.getLogger(__name__)


def remove_existing_cover(doc: DocumentObject) -> None:
    """
    Remove paragraphs that form the existing cover.
    Strategy: delete all paragraphs before the first recognised section heading.
    If no heading is found, nothing is removed (safe fallback).
    """
    paragraphs = list(doc.paragraphs)
    first_heading_index: int | None = None

    for i, p in enumerate(paragraphs):
        if is_section_heading(p):
            first_heading_index = i
            break

    if first_heading_index is None:
        logger.warning("No section heading found — skipping cover removal.")
        return

    body = doc.element.body
    for p in paragraphs[:first_heading_index]:
        body.remove(p._p)

    logger.debug("Removed %d cover paragraphs.", first_heading_index)


def format_body(doc: DocumentObject, cover: CoverData) -> None:
    """Apply ABNT styles to all body paragraphs."""
    font = cover.font.value
    font_size = Pt(12)

    paragraphs = list(doc.paragraphs)
    in_references = False

    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        if is_references_heading(paragraph):
            in_references = True
            _apply_heading_style(paragraph, font, font_size)
            _ensure_page_break_before(paragraphs, i)
            continue

        if in_references:
            _apply_references_style(paragraph, font, font_size)
            continue

        if is_section_heading(paragraph):
            _apply_heading_style(paragraph, font, font_size)
            _ensure_page_break_before(paragraphs, i)
            continue

        _apply_body_style(paragraph, font, font_size)


# ---------------------------------------------------------------------------
# Style appliers
# ---------------------------------------------------------------------------

def _apply_body_style(paragraph: Paragraph, font: str, font_size: object) -> None:
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for run in paragraph.runs:
        apply_run_style(run, font_name=font, font_size=font_size, bold=False)


def _apply_heading_style(paragraph: Paragraph, font: str, font_size: object) -> None:
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for run in paragraph.runs:
        apply_run_style(run, font_name=font, font_size=font_size, bold=True)


def _apply_references_style(paragraph: Paragraph, font: str, font_size: object) -> None:
    """
    References: single spacing, left-aligned, no indent, hanging indent effect
    is achieved via space_after between entries.
    NBR 6023: entries separated by blank line, single spacing within entry.
    """
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.0
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)  # visual separation between entries

    for run in paragraph.runs:
        apply_run_style(run, font_name=font, font_size=font_size, bold=None)


# ---------------------------------------------------------------------------
# Page break helper
# ---------------------------------------------------------------------------

def _ensure_page_break_before(paragraphs: list[Paragraph], index: int) -> None:
    """
    Insert a page break at the END of the paragraph immediately before *index*,
    but only if there isn't one already.
    Skips empty paragraphs when looking for the predecessor.
    """
    predecessor: Paragraph | None = None
    for p in reversed(paragraphs[:index]):
        if p.text.strip():
            predecessor = p
            break

    if predecessor is None or paragraph_has_page_break(predecessor):
        return

    predecessor.add_run().add_break(WD_BREAK.PAGE)
