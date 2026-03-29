from __future__ import annotations

from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.schemas.document import CoverData, TitlePageData
from app.utils.docx_helpers import apply_run_style, is_section_heading

_PAGE_HEIGHT_PT = 841.89
_MARGIN_TOP_PT = 85.04
_MARGIN_BOTTOM_PT = 56.69
_LINE_HEIGHT_PT = 18
_NOTE_INDENT = Cm(8)


def generate_title_page(doc: DocumentObject, cover: CoverData, title_page: TitlePageData) -> None:
    if not title_page.include:
        return

    font = cover.font.value
    font_size = Pt(12)

    note_text = _build_note_text(cover, title_page)
    content_lines = len(cover.authors) + 2 + 3 + 2  # autores + título + nota estimada + cidade/ano
    usable_height = _PAGE_HEIGHT_PT - _MARGIN_TOP_PT - _MARGIN_BOTTOM_PT
    free_space = max(usable_height - content_lines * _LINE_HEIGHT_PT, 0)

    space_before_title = Pt(free_space * 0.42)
    space_before_note = Pt(free_space * 0.08)
    space_before_bottom = Pt(free_space * 0.46)

    paragraphs = []  # (text, alignment, space_before, bold, is_note)

    for author in cover.authors:
        paragraphs.append((author.upper(), WD_ALIGN_PARAGRAPH.CENTER, Pt(0), False, False))

    title_text = cover.title.upper()
    if cover.subtitle:
        subtitle_clean = cover.subtitle.lstrip(": ").strip()
        title_text = f"{title_text}: {subtitle_clean}"

    paragraphs.append((title_text, WD_ALIGN_PARAGRAPH.CENTER, space_before_title, True, False))

    if cover.volume is not None:
        paragraphs.append((f"Volume {cover.volume}", WD_ALIGN_PARAGRAPH.CENTER, Pt(0), False, False))

    # Bloco de natureza como parágrafo único
    paragraphs.append((note_text, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before_note, False, True))

    paragraphs.append((cover.city, WD_ALIGN_PARAGRAPH.CENTER, space_before_bottom, False, False))
    paragraphs.append((str(cover.year), WD_ALIGN_PARAGRAPH.CENTER, Pt(0), False, False))

    insert_position = _find_cover_page_break_index(doc) + 1

    for text, alignment, space_before, bold, is_note in reversed(paragraphs):
        p = _insert_paragraph_at(doc, insert_position)
        p.alignment = alignment
        pf = p.paragraph_format
        pf.space_before = space_before or Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        pf.first_line_indent = Pt(0)
        if is_note:
            pf.left_indent = _NOTE_INDENT
        run = p.add_run(text)
        apply_run_style(run, font_name=font, font_size=font_size, bold=bold)


def remove_existing_title_page(doc: DocumentObject) -> None:
    paragraphs = list(doc.paragraphs)
    heading_indices = [i for i, p in enumerate(paragraphs) if is_section_heading(p)]

    if not heading_indices:
        return

    first_heading = heading_indices[0]
    body = doc.element.body
    page_break_indices = []

    for i, p in enumerate(paragraphs[:first_heading]):
        for el in p._p.iter():
            if el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
                page_break_indices.append(i)
                break

    if len(page_break_indices) < 2:
        return

    start = page_break_indices[0] + 1
    end = page_break_indices[1]
    for p in paragraphs[start:end]:
        body.remove(p._p)


def _build_note_text(cover: CoverData, tp: TitlePageData) -> str:
    """Build the nature note as a single block of text per ABNT NBR 14724."""
    parts = []
    work_type = tp.work_type.value if tp.work_type else "Trabalho acadêmico"

    if tp.work_type and tp.degree and cover.institution:
        parts.append(f"{work_type} apresentado à {cover.institution}, como requisito parcial para obtenção do grau de {tp.degree.value}.")
    elif tp.work_type and cover.institution:
        parts.append(f"{work_type} apresentado à {cover.institution}.")
    else:
        parts.append(f"{work_type}.")

    if tp.concentration_area:
        parts.append(f"Área de concentração: {tp.concentration_area}.")
    if tp.advisor:
        parts.append(f"Orientador: {tp.advisor}.")
    if tp.co_advisor:
        parts.append(f"Coorientador: {tp.co_advisor}.")

    return " ".join(parts)


def _find_cover_page_break_index(doc: DocumentObject) -> int:
    for i, p in enumerate(doc.paragraphs):
        for el in p._p.iter():
            if el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
                return i
    return 0


def _insert_paragraph_at(doc: DocumentObject, position: int):
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    paragraphs = list(doc.paragraphs)
    new_p = OxmlElement("w:p")
    if position < len(paragraphs):
        ref = paragraphs[position]._p
        body.insert(list(body).index(ref), new_p)
    else:
        body.append(new_p)
    return Paragraph(new_p, doc)


def _append_page_break_at(doc: DocumentObject, position: int) -> None:
    body = doc.element.body
    paragraphs = list(doc.paragraphs)
    p_el = OxmlElement("w:p")
    r_el = OxmlElement("w:r")
    br_el = OxmlElement("w:br")
    br_el.set(qn("w:type"), "page")
    r_el.append(br_el)
    p_el.append(r_el)
    if position < len(paragraphs):
        ref = paragraphs[position]._p
        body.insert(list(body).index(ref), p_el)
    else:
        body.append(p_el)