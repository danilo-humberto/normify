from __future__ import annotations

import unicodedata

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor
from docx.styles.style import ParagraphStyle
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run


FONT_NAME = "Arial"
TITLE_FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
TITLE_COLOR = RGBColor(0, 0, 0)
TITLE_KEYWORDS = {"INTRODUCAO", "CONCLUSAO", "REFERENCIAS"}
VALID_SECTIONS = {
    "INTRODUCAO",
    "FUNDAMENTACAO TEORICA",
    "METODOLOGIA",
    "RESULTADOS",
    "DISCUSSAO",
    "CONCLUSAO",
    "REFERENCIAS",
}


def format_docx(input_path: str, output_path: str) -> None:
    document = Document(input_path)

    apply_global_styles(document)

    paragraphs_processed = 0
    references_detected = False
    in_references_section = False
    first_valid_section_found = False

    paragraphs = iter_paragraphs(document)

    for index, paragraph in enumerate(paragraphs):
        paragraphs_processed += 1
        title_detected = is_title(paragraph)
        valid_section_detected = is_valid_section(paragraph)

        if valid_section_detected and should_insert_section_break(
            paragraphs,
            index,
            first_valid_section_found,
        ):
            insert_page_break_before(paragraph)

        if title_detected:
            apply_title_style(paragraph)

        if valid_section_detected:
            first_valid_section_found = True

        if is_references_heading(paragraph):
            references_detected = True
            in_references_section = True
            format_paragraph(paragraph)
            apply_title_style(paragraph)
            continue

        if title_detected:
            continue

        if in_references_section:
            handle_references(paragraph)
            continue

        format_paragraph(paragraph)
        _capitalize_paragraph_runs(paragraph)

    document.save(output_path)

    print(f"Paragraphs processed: {paragraphs_processed}")
    print(f"References section detected: {references_detected}")


def apply_global_styles(document: DocumentObject) -> None:
    for section in document.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = FONT_SIZE
    _set_style_font_family(normal_style, FONT_NAME)


def format_paragraph(paragraph: Paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(1.25)

    for run in paragraph.runs:
        apply_run_font(run, FONT_NAME, bold=False, color=None)


def is_title(paragraph: Paragraph) -> bool:
    text = normalize_text(paragraph.text)
    if not text:
        return False

    is_uppercase = paragraph.text.strip().isupper()
    return is_uppercase or text in TITLE_KEYWORDS


def is_valid_section(paragraph: Paragraph) -> bool:
    return normalize_text(paragraph.text) in VALID_SECTIONS


def apply_title_style(paragraph: Paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Cm(0)

    for run in paragraph.runs:
        apply_run_font(run, TITLE_FONT_NAME, bold=True, color=TITLE_COLOR)


def _capitalize_paragraph_runs(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue

        capitalized_text = capitalize_paragraph(run.text)
        if capitalized_text != run.text:
            run.text = capitalized_text
            return


def insert_page_break_before(paragraph: Paragraph) -> None:
    break_paragraph = paragraph.insert_paragraph_before()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def should_insert_section_break(
    paragraphs: list[Paragraph],
    current_index: int,
    first_valid_section_found: bool,
) -> bool:
    if first_valid_section_found:
        return True

    return has_meaningful_content_before(paragraphs, current_index)


def handle_references(paragraph: Paragraph) -> None:
    format_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = Cm(0)


def is_references_heading(paragraph: Paragraph) -> bool:
    return "REFERENCIAS" in normalize_text(paragraph.text)


def has_meaningful_content_before(paragraphs: list[Paragraph], current_index: int) -> bool:
    for previous_paragraph in paragraphs[:current_index]:
        if previous_paragraph.text.strip():
            return True

    return False


def iter_paragraphs(document: DocumentObject) -> list[Paragraph]:
    paragraphs = list(document.paragraphs)

    for table in document.tables:
        paragraphs.extend(_iter_table_paragraphs(table))

    return paragraphs


def _iter_table_paragraphs(table: Table) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []

    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(_iter_cell_paragraphs(cell))

    return paragraphs


def _iter_cell_paragraphs(cell: _Cell) -> list[Paragraph]:
    paragraphs = list(cell.paragraphs)

    for table in cell.tables:
        paragraphs.extend(_iter_table_paragraphs(table))

    return paragraphs


def apply_run_font(run: Run, font_name: str, bold: bool, color: RGBColor | None) -> None:
    run.font.name = font_name
    run.font.size = FONT_SIZE
    run.font.bold = bold

    if color is not None:
        run.font.color.rgb = color

    if run._element.rPr is None:
        run._element.get_or_add_rPr()

    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
        font_name,
    )
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
        font_name,
    )
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs",
        font_name,
    )


def _set_style_font_family(style: ParagraphStyle, font_name: str) -> None:
    style_rpr = style.element.get_or_add_rPr()
    style_rpr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
        font_name,
    )
    style_rpr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
        font_name,
    )
    style_rpr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs",
        font_name,
    )


def capitalize_paragraph(text: str) -> str:
    for index, character in enumerate(text):
        if character.isalpha():
            return text[:index] + character.upper() + text[index + 1 :]

    return text


def normalize_text(text: str) -> str:
    normalized = text.strip().upper()
    normalized = unicodedata.normalize("NFKD", normalized)
    return normalized.encode("ascii", "ignore").decode("ascii")
